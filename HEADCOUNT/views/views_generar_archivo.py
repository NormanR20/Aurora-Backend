from decimal import MIN_EMIN
from distutils.cygwinccompiler import Mingw32CCompiler
from email import message
from email.policy import HTTP
from unicodedata import decimal
from urllib import response
from django.db.models.functions import ExtractYear
from re import S, sub
from django.contrib.auth.models import User,Group
from django.http.response import Http404
from django.shortcuts import render
from calendar import monthrange
from HEADCOUNT.serializers.serializers_descriptor_perfil import descriptor_perfil_competenciaserializer
from rest_framework.generics import get_object_or_404
from rest_framework.views import APIView
from rest_framework.response import Response
from django.core.mail import EmailMultiAlternatives
from rest_framework import serializers, status
from rest_framework.authentication import SessionAuthentication, BasicAuthentication,TokenAuthentication
from rest_framework.permissions import IsAuthenticated,DjangoModelPermissions
from django.contrib.auth.decorators import login_required, user_passes_test,REDIRECT_FIELD_NAME
from django.utils.decorators import method_decorator
from  rest_framework.authtoken.models import Token
from django.conf import settings
from django.template.loader import render_to_string
from django.utils.html import strip_tags
from django.core.mail import BadHeaderError, send_mail
from django.db.models import *
from django.utils.crypto import get_random_string
import string
from django.contrib.auth import authenticate
from pyrfc import *
from datetime import datetime,timedelta
import json
from ..serializers import *
from ..models import *
from django.apps import apps
from datetime import datetime,timedelta
from django.core.exceptions import ObjectDoesNotExist
import sys
from rest_framework import generics
import io
import csv
sys.setrecursionlimit(100000000)                                                        
from rest_framework import viewsets
import requests
import docx
from docx import Document
from docx.shared import Inches, Cm,Pt,RGBColor
from django.http import HttpResponse
from django.shortcuts import render
from django.http import HttpResponse
from rest_framework import generics
from django.http import HttpResponse
from wsgiref.util import FileWrapper
from docx.enum.table import WD_ROW_HEIGHT_RULE
from django.db.models import CharField, Value
from django.db.models.functions import Concat
from django.conf import settings as django_settings
import base64


class CrearArchivoViewset(generics.ListAPIView):
    authentication_classes=[TokenAuthentication]
    permission_classes=[DjangoModelPermissions]
    queryset = descriptor_perfil_datos_generales.objects.all()
    serializer_class = descriptor_perfil_datos_generalesserializer
    def get(self, request, id, format=None):
        archivo=Document()
        section = archivo.sections[0] 
        header = section.header 
        header_para = header.paragraphs[0] 
        header_para.text = "\t\tF-651V3"

        ##encabezado
        archivo.add_heading("1. DESCRIPTOR y PERFIL DE PUESTOS",1)
        #primer parrafo
        #archivo.add_paragraph('Este es un ejemplo de manula de descriptor', style='Intense Quote')

        descriptor_objeto=descriptor_perfil_datos_generales.objects.get(id=id) if descriptor_perfil_datos_generales.objects.filter(id=id) else None
        #print(descriptor_objeto.empresa.nombre)
        if descriptor_objeto==None:
            return Response({"mensaje":"el descriptor no existe"},status=status.HTTP_404_NOT_FOUND)
        lista_sustitutas=descriptor_objeto.posiciones_sustitutas.all().values()
        sustitutas=descriptor_perfil_datos_generales.objects.filter(id=id).annotate(nombre_sustituta=Concat(F('posiciones_sustitutas__codigo'),Value('-') ,F('posiciones_sustitutas__descripcion')) ) if descriptor_perfil_datos_generales.objects.filter(id=id) else None
        milista_sustituta=''
        if sustitutas:
            for sus in sustitutas:
                if len(milista_sustituta)==0:
                    milista_sustituta =  milista_sustituta+''+sus.nombre_sustituta    
                else:               
                    milista_sustituta =  milista_sustituta+','+sus.nombre_sustituta
        nombre_empresa=str(descriptor_objeto.empresa.nombre) if descriptor_objeto.empresa and descriptor_objeto.empresa.nombre else ''
        num_ocupante=str(descriptor_objeto.numero_ocupantes) if descriptor_objeto.numero_ocupantes else ''
        divi=str(descriptor_objeto.division.descripcion) if descriptor_objeto.division.descripcion else ''
        num_per_sup=str(descriptor_objeto.numero_personas_supervisa) if descriptor_objeto.numero_personas_supervisa else ''
        mmilista_sustituta=milista_sustituta if milista_sustituta else ''
        clasi= str(descriptor_objeto.clasificacion_empleado.nombre) if  descriptor_objeto.clasificacion_empleado else ''
        posi= str(descriptor_objeto.posicion.descripcion) if descriptor_objeto.posicion and descriptor_objeto.posicion.descripcion else ''
        posi_codigo= str(descriptor_objeto.posicion.codigo) if descriptor_objeto.posicion and descriptor_objeto.posicion.codigo else ''
        fecha = str(descriptor_objeto.posicion.nombre)
        uni_orga= str(descriptor_objeto.unidad_organizativa.nombre) if  descriptor_objeto.unidad_organizativa and descriptor_objeto.unidad_organizativa.nombre else ''
        uni_orga_codigo= str(descriptor_objeto.unidad_organizativa.codigo) if  descriptor_objeto.unidad_organizativa and descriptor_objeto.unidad_organizativa.codigo else ''


        records = (
            ("Empresa:", '',nombre_empresa, ''),
            ("Nombre de la posición", '', posi_codigo+'-'+posi, ''),
            ("Unidad Organizativa:", '', uni_orga_codigo+'-'+uni_orga, ''),
            ("Numero de Ocupantes:", '', num_ocupante, ''),
            ("División:", '',divi, ''),
            ("Número de Personas que supervisa:", '', num_per_sup, ''),
            ("Posición que Puede Cubrir las funciones:", '', mmilista_sustituta, ''),
            ("Clasificación:", '',  clasi, ''),
            
        )
        
        #tabla de datos genereales
        datos_generales = archivo.add_table(rows=1, cols=4,style="Table Grid")
        #encabezados de tabla
        hdr_cells = datos_generales.rows[0].cells
        hdr_cells[0].text = 'Dato'
        hdr_cells[2].text = 'Información'
        
        #insertando filas a la tabla
        for dato,espacio1,valor,espacio2  in records:
            row_cells = datos_generales.add_row().cells
            row_cells[0].text = str(dato)
            row_cells[2].text = str(valor)

        for fila in datos_generales.rows:       
            fila.cells[0].merge(fila.cells[1])
            fila.cells[2].merge(fila.cells[3])
        #print(descriptor_objeto.fecha_creacion)
        #print(descriptor_objeto.fecha_actualizacion)
        fecha_ac=str(descriptor_objeto.fecha_actualizacion)
        fecha_ac_fin = fecha_ac
        fecha_cre=str(descriptor_objeto.fecha_creacion_sinhora)
        nueva_fila_datos_generales = datos_generales.add_row().cells
        nueva_fila_datos_generales[0].text = "Fecha Elaboracion"
        nueva_fila_datos_generales[1].text = fecha_cre
        nueva_fila_datos_generales[2].text = "Fecha Actualización"
        nueva_fila_datos_generales[3].text =  str(fecha_ac_fin)



        
        # nueva_fila_datos_generales_ficha = datos_generales_creacion_actualizacion.add_row().cells
        # nueva_fila_datos_generales_ficha[0].text = "ejemplo"
        # nueva_fila_datos_generales_ficha[1].text = "valor1"
        # nueva_fila_datos_generales_ficha[2].text = "ejemplo1"
        # nueva_fila_datos_generales_ficha[3].text = "valor2"

        ##tabla proposito general

        archivo.add_heading("Proposito general del puesto",1)
        #archivo.add_paragraph('Este es el proposito general', style='Intense Quote')
        tabla_proposito_general = archivo.add_table(rows=1, cols=1,style="Table Grid")
        hdr_cells = tabla_proposito_general.rows[0].cells
        hdr_cells[0].text = descriptor_objeto.proposito_general.descripcion if descriptor_objeto.proposito_general and descriptor_objeto.proposito_general.descripcion else ''
        #row_cells = tabla_proposito_general.add_row().cells
        #tabla_proposito_general.rows[1].cells[0].text="proposito general"
        #row_cells = tabla_proposito_general.add_row().cells
        #tabla_proposito_general.rows[1].cells[0].text=descriptor_objeto.proposito_general.descripcion

    
       
        ##tabla propositos
       
        # archivo.add_heading("Propositos del puesto",1)
        tabla_propositos = archivo.add_table(rows=0, cols=1,style="Table Grid")
        #encabezados de tabla
        # hdr_cells = tabla_propositos.rows[0].cells
        # hdr_cells[0].text = 'Propositos'
        prop_obj=descriptor_perfil_proposito_descriptor.objects.filter(descriptor__id=id)
        for x in prop_obj:
            row_cells = tabla_propositos.add_row().cells
            row_cells[0].text = str(x.proposito.nombre)


        

        archivo.add_heading("Funciones del Puesto",1)
        tabla_funciones = archivo.add_table(rows=1, cols=6,style="Table Grid")
        #encabezados de tabla
        hdr_cells = tabla_funciones.rows[0].cells
        hdr_cells[0].width = Cm(8)
        hdr_cells[0].text = 'Funciones'
        hdr_cells[1].width = Cm(1)
        hdr_cells[1].text = 'P'
        hdr_cells[2].width = Cm(1)
        hdr_cells[2].text = 'CE'
        hdr_cells[3].width = Cm(1)
        hdr_cells[3].text = 'CM'
        hdr_cells[4].width = Cm(1.5)
        hdr_cells[4].text = 'Total'
        hdr_cells[5].width = Cm(3)
        hdr_cells[5].text = 'Fundamentales'

        
        lista_funciones=(
            (1,2,3,4,5,6),
            (1,2,3,4,5,6)

        )
        lista_funciones=list(descriptor_perfil_funcion.objects.filter(descriptor__id=id).values_list('descripcion','perioricidad','concecuencias_error','complejidad','fundamental'))
        



        #insertando filas a la tabla
        for Funciones,P, CE,CM, Fundamentales in lista_funciones:
            row_cells = tabla_funciones.add_row().cells
            total=float(str(P) if P else 0) + ((float(str(CE)) if CE else 0) * (float(str(CM)) if CM else 0) )
            row_cells[0].text = str(Funciones)
            row_cells[1].text = str(P)
            row_cells[2].text = str(CE)
            row_cells[3].text = str(CM)
            row_cells[4].text = str(total)
            row_cells[5].text = 'Si' if Fundamentales else 'No'

        ##indicadores
        archivo.add_heading("Indicadores:",1)
        indicadores = archivo.add_table(rows=1, cols=5,style="Table Grid")
        #encabezados de tabla
        hdr_cells = indicadores.rows[0].cells
        hdr_cells[0].text = 'Objetivo'
        hdr_cells[1].text = 'Nombre'
        hdr_cells[2].text = 'Formula'
        hdr_cells[3].text = 'Meta'
        hdr_cells[4].text = 'Fuente Verificacion'
        

        lista_indicadores_resultado=list(descriptor_perfil_indicador_descriptor.objects.filter(descriptor__id=id).values_list('indicador__objetivo','indicador__nombre','indicador__formula_calculo','meta','fuente_verificacion'))


        
        #insertando filas a la tabla
        for objetivo,nombre, formula,meta,fuente_verificacion in lista_indicadores_resultado:
            row_cells = indicadores.add_row().cells
            row_cells[0].text = str(objetivo)
            row_cells[1].text = str(nombre)
            row_cells[2].text = str(formula)
            row_cells[3].text = str(meta)
            row_cells[4].text = str(fuente_verificacion)
            
            
        ##matriz competencias
        archivo.add_heading("Matriz de Competencias:",1)
        
        tabla_competencia = archivo.add_table(rows=1, cols=3,style="Table Grid")
        #encabezados de tabla
        hdr_cells = tabla_competencia.rows[0].cells
        hdr_cells[0].text = 'Tipo Competencia'
        hdr_cells[1].text = 'Competencia'
        hdr_cells[2].text = 'Nivel de Desarrollo'


        lista_competencia=descriptor_perfil_competencia_descriptor.objects.filter(descriptor__id=id).values_list('tipo_competencia__descripcion','nombre','nivel_desarrollo').order_by('tipo_competencia')
        #insertando filas a la tabla

        for tipo_competencia,competencia, nivel_desarrollo in lista_competencia:
            row_cells = tabla_competencia.add_row().cells
            row_cells[0].text = str(tipo_competencia)
            row_cells[1].text = str(competencia)
            row_cells[2].text = str(nivel_desarrollo)

        for rw in tabla_competencia.rows:
        
            for rw2 in tabla_competencia.rows:
                if rw.cells[0].text ==rw2.cells[0].text and rw._index!=rw2._index:
                    tabla_competencia.cell(rw._index,0).merge(tabla_competencia.cell(rw2._index,0))



        ##Politicas y procedimientos
        archivo.add_heading("Politicas y Procedimientos:",1)
        politica_procedimiento_tabla = archivo.add_table(rows=0, cols=1,style="Table Grid")
        #encabezados de tabla
        row_cells = politica_procedimiento_tabla.add_row().cells
        hdr_cells = politica_procedimiento_tabla.rows[0].cells
        hdr_cells[0].text = 'Politica Y Procedimiento'
 
        lista_politica_procedimiento=list(descriptor_perfil_politicas_procedimientos.objects.filter(descriptor__id=id).values_list('nombre',flat=True))
        #print("lista",lista_politica_procedimiento)
        #insertando filas a la tabla
        for pp in lista_politica_procedimiento:
            row_cells = politica_procedimiento_tabla.add_row().cells
            row_cells[0].text = str(pp)


        archivo.add_heading("Informacíon Perfil del Puesto:",1)
        perfil = archivo.add_table(rows=1, cols=8,style="Table Grid")
        #encabezados de tabla
        hdr_cells = perfil.rows[0].cells
        hdr_cells[0].text = 'Edad'
        hdr_cells[3].text = 'Género'
        hdr_cells[6].text = 'Viajes de Trabajo'

        row_cells = perfil.add_row().cells

        perfil.rows[0].cells[0].merge(perfil.rows[0].cells[1])

        perfil.rows[0].cells[2].merge(perfil.rows[0].cells[3]).merge(perfil.rows[0].cells[4])
        perfil.rows[0].cells[5].merge(perfil.rows[0].cells[6]).merge(perfil.rows[0].cells[7])
        hdr_cells = perfil.rows[1].cells
        hdr_cells[0].text = 'De'
        hdr_cells[1].text = 'A'
        hdr_cells[2].text = 'Femenino'
        hdr_cells[3].text = 'Masculino'
        hdr_cells[4].text = 'Indiferente'
        hdr_cells[5].text = 'SI'
        hdr_cells[6].text = 'NO'
        hdr_cells[7].text = 'Ocacionalmente'

        perfil.add_row()
        hdr_cells = perfil.rows[2].cells
        hdr_cells[0].text = str(descriptor_objeto.edad_minima)
        hdr_cells[1].text = str(descriptor_objeto.edad_maxima)
        hdr_cells[2].text = str('X' if descriptor_objeto.genero =='Femenino' else '')
        hdr_cells[3].text = str('X' if descriptor_objeto.genero =='Masculino' else '')
        hdr_cells[4].text = str('X' if descriptor_objeto.genero =='Indiferente' else '')
        hdr_cells[5].text = str('X' if descriptor_objeto.necesita_viajar=='Si' else '')
        hdr_cells[6].text = str('X' if descriptor_objeto.necesita_viajar=='No' else '')
        hdr_cells[7].text = str('X' if descriptor_objeto.necesita_viajar=='Ocacionalmente' else '')


        archivo.add_heading("Conocimiento y Experiencia:",1)
        archivo.add_heading("a)  Educacion Formal Requerida:",3)
        lista_educacion=list(descriptor_perfil_formacion.objects.filter(descriptor__id=id).values_list('formacion__nivel_academico__nombre','formacion__titulo','formacion__area_conocimiento__nombre','indispensable'))
        Educacion_formal = archivo.add_table(rows=1, cols=4,style="Table Grid")
        #encabezados de tabla
        hdr_cells = Educacion_formal.rows[0].cells
        hdr_cells[0].text = 'Nivel de Educación Formal'
        hdr_cells[1].text = 'Titulos requeridos'
        hdr_cells[2].text = 'Areá de conocimientos formales'
        hdr_cells[3].text = 'Deseable o indispensable'


        #row_cells = Educacion_formal.add_row().cells

        for nivel,titulo, area,deseable in lista_educacion:
            row_cells = Educacion_formal.add_row().cells
            row_cells[0].text = str(nivel)
            row_cells[1].text = str(titulo)
            row_cells[2].text = str(area)
            row_cells[3].text = 'Indispensable' if deseable==True else 'Deseable'



        archivo.add_heading("b)  Capacitación adicional requerida:",3)
        capacitacion_adicional = archivo.add_table(rows=1, cols=3,style="Table Grid")
        #encabezados de tabla
        hdr_cells = capacitacion_adicional.rows[0].cells
        hdr_cells[0].text = 'Curso/Seminario/Pasantia'
        hdr_cells[1].text = 'Duracion en Horas'
        hdr_cells[2].text = 'Deseable o Indispensable'

        lista_educacion_capa=(
        (1,2,3),
        (1,2,3),
        (2,3,5),
        (5,6,5),
        (3,6,4),
        (3,1,4)
        )
        lista_educacion_capa=list(descriptor_perfil_preparacion.objects.filter(descriptor__id=id).values_list('curso__nombre','curso__duracion','indispensable'))
        for curso,duracion, deseable in lista_educacion_capa:
            row_cells = capacitacion_adicional.add_row().cells
            row_cells[0].text = str(curso)
            row_cells[1].text = str(duracion)
            row_cells[2].text = 'Indispensable' if deseable==True else 'Deseable'


        archivo.add_heading("c)  Conocimientos Técnicos Especializados:",3)
        conocimiento_tecnico = archivo.add_table(rows=1, cols=4,style="Table Grid")
        #encabezados de tabla
        hdr_cells = conocimiento_tecnico.rows[0].cells
        hdr_cells[0].text = 'Conocimientos Técnicos'
        hdr_cells[1].text = 'Descripción'
        hdr_cells[2].text = 'Nivel de Profundidad'
        hdr_cells[3].text = 'Requerimiento de selección o de Capacitación'

        lista_educacion_tecnico=(
        (1,2,3,4),
        (1,2,3,4),
        (2,3,5,6),
        (5,6,5,7),
        (3,6,4,5),
        (3,1,4,6)
        )
        lista_educacion_tecnico=list(descriptor_perfil_conocimiento_tecnico_adquirido.objects.filter(descriptor__id=id).values_list('nombre','descripcion','nivel_profundidad','indispensable'))

        for conocimiento,descripcion, nivel,requerimiento in lista_educacion_tecnico:
            row_cells = conocimiento_tecnico.add_row().cells
            row_cells[0].text = str(conocimiento)
            row_cells[1].text = str(descripcion)
            
            if str(nivel)=='1':
                row_cells[2].text = 'Básico' 
            elif str(nivel) == '2':
                row_cells[2].text = 'Intermedio'
            elif str(nivel) == '3':
                row_cells[2].text = 'Avanzado'
            else:
                row_cells[2].text = ''
                
            row_cells[3].text = 'indispensable' if requerimiento else 'deseable'



        archivo.add_heading("d) Experiencia laboral requerida:",3)
        experiencia_requerida = archivo.add_table(rows=1, cols=5,style="Table Grid")
        #encabezados de tabla
        hdr_cells = experiencia_requerida.rows[0].cells
        hdr_cells[0].text = 'Dimensión Experiencia'
        hdr_cells[1].text = 'Experiencia Instituciones'
        hdr_cells[3].text = 'Experiencia en Cargos'

     
        hdr_cells[1].merge(hdr_cells[1]).merge(hdr_cells[2])
        hdr_cells[3].merge(hdr_cells[3]).merge(hdr_cells[4])
        # row_cells = experiencia_requerida.add_row().cells
        # row_cells[0].text = "Tiempo"
        # row_cells[1].text = "Exp. Institucion tipo"
        # row_cells[2].text = "EXP. Puesto tipo"
        # row_cells[3].text = "Institucion"
        # row_cells[4].text = "Detalle"


        lista_experiencia_requerida=list(descriptor_perfil_experiencia.objects.filter(descriptor__id=id).values_list('tiempo','experiencia_tipo_institucion','experiencia_tipo_cargo','institucion','experiencia_tipo_institucion_detalle','tiempo_2','experiencia_tipo_cargo_detalle'))

        for tiempo,instipo,puestotipo,institucion,institucion_detalle,tiempo_2,tipo_cargo_detalle in lista_experiencia_requerida:
            row_cells = experiencia_requerida.add_row().cells
            # if tiempo==None:
            #     tiempo='Tiempo no ingresado'
            # else:
            #     tiempo= int(tiempo)

            # if tiempo_2==None:
            #     tiempo_2='Tiempo no ingresado'
            # else:
            #     tiempo_2= int(tiempo_2)

            if tiempo!=None and tiempo_2!=None:

                row_cells[0].text = str( str(int(tiempo)) +' a '+ str(int(tiempo_2))+' años' )
            else:
                 row_cells[0].text = str('Tiempos no ingresados')
            row_cells[1].text = 'Si' if instipo ==True else 'No'
            row_cells[2].text = str(institucion_detalle)
            row_cells[3].text = 'Si' if puestotipo==True else 'No'
            row_cells[4].text = str(tipo_cargo_detalle)
        archivo.add_page_break()

        archivo.add_heading("OBSERVACIONES:",1)
        archivo.add_paragraph('Utilice este espacio en caso de querer señalar alguna observación, sugerencia, comentario o acuerdos de re editar o actualizar esta descripción.',style = 'Subtitle')
        observacion = archivo.add_table(rows=1, cols=1, style='Table Grid')
        observacion.autofit = False 
        observacion.allow_autofit = False
        hdr_cells = observacion.rows[0].cells
   
          
        # observacion.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        # observacion.rows[1].cells[0].height = Inches(100)
        # observacion.rows[0].cells[0].height = Inches(100)

        # observacion.cell(0,0).width = Inches(1.0)   
        # observacion.cell(1,0).width = Inches(1.0)     
        # observacion.rows[1].height = Inches(2.0)                                            
        #for fila in observacion.rows:
        #observacion.cell(0, 0).add_paragraph().add_run("text here")
        #observacion.cell(0, 0).paragraphs[0].runs[1].font.size = Pt(9)
        

        for tableRow in observacion.rows:
            for cell in tableRow.cells:
                para = cell.paragraphs[0]
                run = para.add_run( ". " )
                font = run.font
                font = run.font
                font.color.rgb = RGBColor(255, 255, 255)
                font.size = Pt( 60 )
      
        archivo.add_heading("AUTORIZACIONES:",1)
        archivo.add_heading("Fecha:",4)

 
        
        
        archivo.add_heading("___________________        _______________       _________________",1)
        firmas=archivo.add_heading("OCUPANTE DEL PUESTO                                        JEFE INMEDIATO                                        RECURSOS HUMANOS",7)
        run = firmas.runs[0]
        font = run.font
        font.name = 'Calibri'
        font.size = Pt(7)
        nombres=archivo.add_heading("Nombre:                                                                   Nombre:                                                   Nombre:",7)

        run = nombres.runs[0]
        font = run.font
        font.name = 'Calibri'
        font.size = Pt(7)
        
        nombre_archivo_creado=str(settings.STATICFILES_DIRS[0]) + '/descriptores/'
        nombre_archivo_creado=nombre_archivo_creado+request.user.username +descriptor_objeto.posicion.descripcion+'.docx'
        #print(nombre_archivo_creado)
        
        archivo.save(nombre_archivo_creado)
        #documento = open(nombre_archivo_creado, 'rb')
        archivoword_codificado=''
        with open(nombre_archivo_creado, "rb") as archivoword:
            archivoword_codificado = base64.b64encode(archivoword.read()).decode('utf-8')
            
        #print('archivo base64',archivoword_codificado)
        #print('archivo string',str(archivoword_codificado))
        # response = HttpResponse(FileWrapper(documento), content_type='application/msword')
        # response['Content-Disposition'] = 'attachment; filename="%s"' % descriptor_objeto.nombre_posicion+'.docx'
        return Response({"documento":str(archivoword_codificado)},status= status.HTTP_200_OK)
    

# import requests
# def prueba():
#     url = "http://127.0.0.1:8000/api/app/Descargar_Manual_Descriptor/2/"

#     payload={}
#     headers = {
#     'Authorization': 'token 0e7747c279410ddb6acc0ae1fba307c9b0536d92',
#     'Content-Type': 'application/json'
#     }

#     response = requests.get( url, headers=headers, data=payload, stream=True)
#     with open("prueba3.docx", 'wb') as f:
#       for chunk in response.iter_content(1024 * 1024 * 2):  # 2 MB chunks
#         f.write(chunk)


#     print(response.text)



class seleccion_contratacion_plan_requerimientoViewset(generics.ListAPIView):
    authentication_classes=[TokenAuthentication]
    permission_classes=[DjangoModelPermissions]
    queryset = descriptor_perfil_datos_generales.objects.all()
    serializer_class = descriptor_perfil_datos_generalesserializer
    def get(self, request, id, format=None):
        archivo=Document()
        # archivo.add_paragraph('F-329 V3\tRight')
        section = archivo.sections[0]
        header = section.header
        unicode_character = "\u2611"
        paragraph = header.paragraphs[0]
        paragraph.text = "\t\tF-329 V3"
        paragraph.style = archivo.styles["Header"]
        styles = archivo.styles
        styles['Heading 1'].font.color.rgb = docx.shared.RGBColor(0, 0, 0) #establece el color del titulo
        titulo=archivo.add_heading("REQUISICION DE PERSONAL",1)
        parrafo = archivo.add_paragraph(" ")
        titulo.alignment = 1
        
        solicitud_objeto = seleccion_contratacion_solicitud_plaza_vacante.objects.get(id=id) if seleccion_contratacion_solicitud_plaza_vacante.objects.filter(id=id) else None 
        if solicitud_objeto==None:
            return Response({"mensaje":"la solicittud no existe"},status=status.HTTP_404_NOT_FOUND)

        jefe_codigo = str(solicitud_objeto.creador_plaza.username) if solicitud_objeto.creador_plaza else ''
        if jefe_codigo!=None:
            nombre_jefe=(Funcional_empleado.objects.filter(codigo=jefe_codigo).values('nombre'))[0]['nombre'] if Funcional_empleado.objects.filter(codigo=jefe_codigo).values('nombre') else None
        else:
            nombre_jefe=''
        
        if jefe_codigo!=None:
            funcion_jefe=(Funcional_empleado.objects.filter(codigo=jefe_codigo).values('posicion__descripcion'))[0]['posicion__descripcion'] if Funcional_empleado.objects.filter(codigo=jefe_codigo).values('posicion__descripcion') else None
        else:
            funcion_jefe=''
            
        departamento = str(solicitud_objeto.departamento.descripcion) if solicitud_objeto.departamento else None
        puesto=str(solicitud_objeto.posicion.descripcion) if solicitud_objeto.posicion else None 
        funcion= str(solicitud_objeto.funcion.descripcion) if solicitud_objeto.funcion else None
        funcion_id= str(solicitud_objeto.funcion.id) if solicitud_objeto.funcion else None
        motivo=str(solicitud_objeto.motivo.descripcion) if solicitud_objeto.motivo else None
        numero_ocupantes_valor=str(solicitud_objeto.numero_posiciones if solicitud_objeto and solicitud_objeto.numero_posiciones else '-----')
        fecha_solicitud_valor= str(solicitud_objeto.fecha_solicitud if solicitud_objeto and solicitud_objeto.fecha_solicitud else '------' )
        # print(solicitud_objeto.posicion.descripcion)
        records = (
            ("Jefe:", '',nombre_jefe, ''),
            ("Cargo del Jefe:", '', funcion_jefe, ''),
            ("Departamento:", '', departamento if departamento else '', ''),
            ("Solicitud al Departamento de Recursos Humanos cubrir la vacante de:",'',funcion if funcion else '','')   
            )

         #tabla de datos genereales
        datos_principales = archivo.add_table(rows=0, cols=4,style="Table Grid")
        # encabezados de tabla
        # hdr_cells = datos_principales.rows[0].cells
        # hdr_cells[0].text = 'Dato'
        # hdr_cells[2].text = 'Información'
        
        #insertando filas a la tabla
        for dato,espacio1,valor,espacio2  in records:
            row_cells = datos_principales.add_row().cells
            row_cells[0].text = str(dato)
            row_cells[2].text = str(valor)

        for fila in datos_principales.rows:       
            fila.cells[0].merge(fila.cells[1])
            fila.cells[2].merge(fila.cells[3])
        
        parrafo = archivo.add_paragraph(" ")
        # parrafo.add_run().add_break()

        if motivo=='Sustitución':
            p = archivo.add_paragraph('Por:')
            a=archivo.add_paragraph('Sustitución       (',style='List Bullet')
            b=archivo.add_paragraph('Nueva vacante (  )',style='List Bullet')
            c=archivo.add_paragraph('Otro                  (  )',style='List Bullet')
            paragraph_format_a = a.paragraph_format
            paragraph_format_b = b.paragraph_format
            paragraph_format_c = c.paragraph_format
            a.add_run('X').bold = True
            a.add_run(')')
            paragraph_format_a.left_indent = Inches(0.5)
            paragraph_format_b.left_indent = Inches(0.5)
            paragraph_format_c.left_indent = Inches(0.5)

        if motivo=='Nueva vacante':
            p = archivo.add_paragraph('Por:')
            a=archivo.add_paragraph('Sustitución       (  )',style='List Bullet')
            b=archivo.add_paragraph('Nueva vacante (',style='List Bullet')
            c=archivo.add_paragraph('Otro                  (  )',style='List Bullet')
            paragraph_format_a = a.paragraph_format
            paragraph_format_b = b.paragraph_format
            paragraph_format_c = c.paragraph_format
            b.add_run('X').bold = True
            b.add_run(')')
            paragraph_format_a.left_indent = Inches(0.5)
            paragraph_format_b.left_indent = Inches(0.5)
            paragraph_format_c.left_indent = Inches(0.5)

        if motivo=='Otro':
            p = archivo.add_paragraph('Por:')
            a=archivo.add_paragraph('Sustitución       (  )',style='List Bullet')
            b=archivo.add_paragraph('Nueva vacante (  )',style='List Bullet')
            c=archivo.add_paragraph('Otro                  (',style='List Bullet')
            paragraph_format_a = a.paragraph_format
            paragraph_format_b = b.paragraph_format
            paragraph_format_c = c.paragraph_format
            c.add_run('X').bold = True
            c.add_run(')')
            paragraph_format_a.left_indent = Inches(0.5)
            paragraph_format_b.left_indent = Inches(0.5)
            paragraph_format_c.left_indent = Inches(0.5)

        styles['Heading 2'].font.color.rgb = docx.shared.RGBColor(0, 0, 0) #establece el color del titulo
        parrafo = archivo.add_paragraph(" ")
        perfil_puesto=archivo.add_heading("Perfil del Puesto",2)
        
        perfil = archivo.add_table(rows=1, cols=10,style="Table Grid")
        #encabezados de tabla
        hdr_cells = perfil.rows[0].cells
        hdr_cells[0].text = 'Edad'
        hdr_cells[3].text = 'Género'
        hdr_cells[6].text = 'Viajes de Trabajo'
        hdr_cells[9].text = 'Vehiculo'

        row_cells = perfil.add_row().cells

        perfil.rows[0].cells[0].merge(perfil.rows[0].cells[1])
        
        perfil.rows[0].cells[2].merge(perfil.rows[0].cells[3]).merge(perfil.rows[0].cells[4])
        perfil.rows[0].cells[5].merge(perfil.rows[0].cells[6]).merge(perfil.rows[0].cells[7])
        perfil.rows[0].cells[8].merge(perfil.rows[0].cells[9])
        hdr_cells = perfil.rows[1].cells
        hdr_cells[0].text = 'De'
        hdr_cells[1].text = 'A'
        hdr_cells[2].text = 'F'
        hdr_cells[3].text = 'M'
        hdr_cells[4].text = 'indiferente'
        hdr_cells[5].text = 'SI'
        hdr_cells[6].text = 'NO'
        hdr_cells[7].text = 'Ocacionalmente'
        hdr_cells[8].text = 'SI'
        hdr_cells[9].text = 'NO'

        for row in perfil.rows:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        font = run.font
                        font.size= Pt(12)

        descriptor_objeto = descriptor_perfil_datos_generales.objects.filter(posicion=funcion_id).order_by('-id') if descriptor_perfil_datos_generales.objects.filter(posicion=funcion_id) else None
        # print('descriptor_objeto',descriptor_objeto)
        descriptor_id = descriptor_objeto[0].id if descriptor_perfil_datos_generales.objects.filter(posicion=funcion_id) else None
        # print('funcion_id',funcion_id)
        # print('descriptor',descriptor_objeto[0].id)
        perfil.add_row()
        hdr_cells = perfil.rows[2].cells
        hdr_cells[0].text = str(descriptor_objeto[0].edad_minima if descriptor_objeto and descriptor_objeto[0].edad_minima else '') 
        hdr_cells[1].text = str(descriptor_objeto[0].edad_maxima if descriptor_objeto and descriptor_objeto[0].edad_maxima else '') 
        hdr_cells[2].text = str('X' if descriptor_objeto and descriptor_objeto[0].genero =='Femenino' else '')
        hdr_cells[3].text = str('X' if descriptor_objeto and descriptor_objeto[0].genero =='Masculino' else '')
        hdr_cells[4].text = str('X' if descriptor_objeto and descriptor_objeto[0].genero =='Indiferente' else '')
        hdr_cells[5].text = str('X' if descriptor_objeto and descriptor_objeto[0].necesita_viajar=='Si' else '')
        hdr_cells[6].text = str('X' if descriptor_objeto and descriptor_objeto[0].necesita_viajar=='No' else '')
        hdr_cells[7].text = str('X' if descriptor_objeto and descriptor_objeto[0].necesita_viajar=='Ocasionalmente' else '')
        hdr_cells[8].text = str('X' if descriptor_objeto and descriptor_objeto[0].vehiculo=='Si' else '')
        hdr_cells[9].text = str('X' if descriptor_objeto and descriptor_objeto[0].vehiculo=='No' else '')
        parrafo = archivo.add_paragraph(" ")
        

        lista_educacion=list(descriptor_perfil_formacion.objects.filter(descriptor__id=descriptor_id).values_list('formacion__nivel_academico__nombre','formacion__titulo','formacion__area_conocimiento__nombre','indispensable')) if descriptor_perfil_formacion.objects.filter(descriptor__id=descriptor_id) else None
        # print('lista_educacion',descriptor_perfil_formacion.objects.filter(descriptor__id=descriptor_id).values_list('id',flat=True))
        Educacion_formal = archivo.add_table(rows=1, cols=4,style="Table Grid")
        #encabezados de tabla
        hdr_cells = Educacion_formal.rows[0].cells
        hdr_cells[0].text = 'Nivel de Educación Formal'
        hdr_cells[1].text = 'Titulos requeridos'
        hdr_cells[2].text = 'Areá de conocimientos formales'
        hdr_cells[3].text = 'Deseable o indispensable'

        if lista_educacion!=None:
            for nivel,titulo, area,deseable in lista_educacion:
                row_cells = Educacion_formal.add_row().cells
                row_cells[0].text = str(nivel)
                row_cells[1].text = str(titulo)
                row_cells[2].text = str(area)
                row_cells[3].text = 'Indispensable' if deseable==True else 'Deseable'
                parrafo = archivo.add_paragraph(" ")

    #    experiencia 
        experiencia= descriptor_perfil_experiencia.objects.filter(descriptor=descriptor_id) if descriptor_perfil_experiencia.objects.filter(descriptor=descriptor_id) else None
        if experiencia!=None:
            records2 = (
            ("1. Tiempo de experiencia:", '',experiencia[0].tiempo, ''),
            ("2. Experiencia en Instituciones de tipo:", '', str('Si' if experiencia[0].experiencia_tipo_institucion==True else 'No'), ''),
            ("3. Experiencia en Cargos de tipo:",'', str('Si' if experiencia[0].experiencia_tipo_cargo==True else 'No'),'')    
             )
        else:
            records2 = (
            ("1. Tiempo de experiencia:", '','Sin Experiencia', ''),
            ("2. Experiencia en Instituciones de tipo:", '', 'No', ''),
            ("3. Experiencia en Cargos de tipo:",'', 'No','')    
            )

        #tabla de datos genereales
        datos_principales_experiencia = archivo.add_table(rows=0, cols=4,style="Table Grid")
     
        for dato,espacio1,valor,espacio2  in records2:
            row_cells = datos_principales_experiencia.add_row().cells
            row_cells[0].text = str(dato)
            row_cells[2].text = str(valor)

        for fila in datos_principales_experiencia.rows:       
            fila.cells[0].merge(fila.cells[1])
            fila.cells[2].merge(fila.cells[3])
        parrafo = archivo.add_paragraph(" ")
        # parrafo = archivo.add_paragraph(" ")
        # parrafo = archivo.add_paragraph(" ")
        #Otros conocimientos 
        
        
        conocimientos = archivo.add_table(rows=1, cols=6,style="Table Grid")
        #encabezados de tabla
        hdr_cells = conocimientos.rows[0].cells
        hdr_cells[0].text = 'Otros Conocimientos o Habilidades (Describirlos)'
        hdr_cells[3].text = '\tNivel de Desarrollo'

        row_cells = conocimientos.add_row().cells
        conocimientos.cell(0, 0).merge(conocimientos.cell(1, 2))

        conocimientos.rows[0].cells[0].merge(conocimientos.rows[0].cells[1]).merge(conocimientos.rows[0].cells[2])
        conocimientos.rows[0].cells[3].merge(conocimientos.rows[0].cells[4]).merge(conocimientos.rows[0].cells[5])

        hdr_cells_conocimientos = conocimientos.rows[1].cells
        hdr_cells_conocimientos[3].text = 'Básico'
        hdr_cells_conocimientos[4].text = 'Intermedio'
        hdr_cells_conocimientos[5].text = 'Avanzado'


        conocimientos_adquiridos=list(descriptor_perfil_conocimiento_tecnico_adquirido.objects.filter(descriptor__id=descriptor_id).values_list('nombre','nivel_profundidad')) if descriptor_perfil_conocimiento_tecnico_adquirido.objects.filter(descriptor__id=descriptor_id) else None
        
        # lista_funciones=list(descriptor_perfil_funcion.objects.filter(descriptor__id=id).values_list('descripcion','perioricidad','concecuencias_error','complejidad','fundamental'))

        #insertando filas a la tabla
        # print('conocimientos_adquiridos',conocimientos_adquiridos)
        if conocimientos_adquiridos!=None:
            for nombre,nivel_profundidad in conocimientos_adquiridos:
                row_cells_c = conocimientos.add_row().cells
                row_cells_c[0].text = str(nombre)
                row_cells_c[3].text = str('\tX' if nivel_profundidad==1 or nivel_profundidad==2 else '')
                row_cells_c[4].text = str('\tX' if nivel_profundidad==3 or nivel_profundidad==4 else '')
                row_cells_c[5].text = str('\tX' if nivel_profundidad==5 else '')
            for fila in conocimientos.rows:       
                fila.cells[0].merge(fila.cells[1]).merge(fila.cells[2])
            
        parrafo = archivo.add_paragraph(" ")
        parrafo = archivo.add_paragraph(" ")
        parrafo = archivo.add_paragraph("Número de Personas que solicita: "+ numero_ocupantes_valor +"       "" Fecha:  "+ fecha_solicitud_valor + "")
        parrafo = archivo.add_paragraph(" ")
        parrafo = archivo.add_paragraph("Entrego Plan de Inducción:               Si "+unicode_character +"                 No ☐")
        parrafo = archivo.add_paragraph(" ")
        parrafo = archivo.add_paragraph("                                                ____________________________                                     ")
        parrafo = archivo.add_paragraph("                                                          Firma                                     ")

        ####################################################################################################################################################
        nombre_archivo_creado=str(settings.STATICFILES_DIRS[0]) + '/seleccion_contratacion/'

        nombre_archivo_creado=nombre_archivo_creado+request.user.username +funcion+'.docx'
      
        
        archivo.save(nombre_archivo_creado)
        archivoword_codificado=''
        with open(nombre_archivo_creado, "rb") as archivoword:
            archivoword_codificado = base64.b64encode(archivoword.read()).decode('utf-8')
            
        
        return Response({"documento":str(archivoword_codificado)},status= status.HTTP_200_OK)




def empleado(codigo,campo):
   resultado= Funcional_empleado.objects.filter(codigo=codigo).values(str(campo))
   print(resultado)
   return resultado[0][campo]


class capacitacion_archivo_evaluacionViewset(generics.ListAPIView):
    authentication_classes=[TokenAuthentication]
    permission_classes=[DjangoModelPermissions]
    queryset = capacitacion_evento_capacitacion.objects.all()
    serializer_class = capacitacion_evento_capacitacionserializer
    def get(self, request, id, format=None):
        archivo=Document()
        # archivo.add_paragraph('F-329 V3\tRight')
        section = archivo.sections[0]
        header = section.header
        unicode_character = "\u2611"
        paragraph = header.paragraphs[0]
        hoy=datetime.now()
        hoy_formateado=  datetime.strftime(hoy, '%d-%m-%Y')
        paragraph.text = "\t\tF-399v4"
        paragraph.style = archivo.styles["Header"]
        styles = archivo.styles
        styles['Heading 1'].font.color.rgb = docx.shared.RGBColor(0, 0, 0) #establece el color del titulo
        titulo=archivo.add_heading("CONVOCATORIA PARA CAPACITACIÓN",1)
        parrafo = archivo.add_paragraph(" ")
        titulo.alignment = 1

        solicitud_objeto = capacitacion_evento_capacitacion.objects.get(id=id) if capacitacion_evento_capacitacion.objects.filter(id=id) else None 
        if solicitud_objeto==None:
            return Response({"mensaje":"El evento no existe"},status=status.HTTP_404_NOT_FOUND)

        mes_evento= (capacitacion_evento_capacitacion.objects.filter(id=id).values('fecha_inicio__month'))[0]['fecha_inicio__month'] if capacitacion_evento_capacitacion.objects.filter(id=id).values('fecha_inicio__month')  else None
        mes_evento_digito= (capacitacion_evento_capacitacion.objects.filter(id=id).values('fecha_inicio__month'))[0]['fecha_inicio__month'] if capacitacion_evento_capacitacion.objects.filter(id=id).values('fecha_inicio__month')  else None
        anio_evento= (capacitacion_evento_capacitacion.objects.filter(id=id).values('fecha_inicio__year'))[0]['fecha_inicio__year'] if capacitacion_evento_capacitacion.objects.filter(id=id).values('fecha_inicio__year')  else None
        dia_evento= (capacitacion_evento_capacitacion.objects.filter(id=id).values('fecha_inicio__day'))[0]['fecha_inicio__day'] if capacitacion_evento_capacitacion.objects.filter(id=id).values('fecha_inicio__day')  else None
        meses= ['Enero','Febrero','Marzo','Abril','Mayo','Junio','Julio','Agosto','Septiembre','Octubre','Noviembre','Diciembre']
        campania= solicitud_objeto.campania.nombre_campania  if solicitud_objeto!=None else None
        impartido_por= solicitud_objeto.impartido_por if solicitud_objeto!=None else None
        if mes_evento!=None:
            indice=(mes_evento)-1
            mes_evento=meses[indice]

        archivo.add_heading(str(dia_evento)+'-'+str(mes_evento_digito)+'-'+str(anio_evento),1)
        archivo.add_heading('Estimado/a Colaborador',1)
        parrafo = archivo.add_paragraph(" ")
        parrafo = archivo.add_paragraph("Nos complace informarle que han sido seleccionados para participar en la Capacitación "+str(campania)+' impartido por '+str(impartido_por) +', la cual se realizara: ' )
        archivo.add_heading('Bajo la Modalidad:',1)
        
        
        modalidades= capacitacion_modalidad.objects.filter().values_list('descripcion',flat=True)
        modalidad_evento=solicitud_objeto.modalidad.descripcion
        
        for modalidad in modalidades:
            # print('modalidad',modalidad)
            parrafo=archivo.add_paragraph(str(modalidad) +'      '+ str(unicode_character if modalidad==modalidad_evento else "☐"))
        parrafo = archivo.add_paragraph(" ")

        parrafo = archivo.add_paragraph('Día/Mes/Año: '+str(dia_evento)+'/'+str(mes_evento_digito)+'/'+str(anio_evento) )

        
        parrafo = archivo.add_paragraph("Listado del o los convocados")
        ##indicadores
        # archivo.add_heading("Indicadores:",1)
        convocados = archivo.add_table(rows=1, cols=4,style="Table Grid")
        #encabezados de tabla
        hdr_cells = convocados.rows[0].cells
        hdr_cells[0].text = 'Nro.'
        hdr_cells[1].text = 'CODIGO'
        hdr_cells[2].text = 'NOMBRE'
        hdr_cells[3].text = 'DEPARTAMENTO'
        
        listado_convocados_capacitacion=list(capacitacion_evento_capacitacion.objects.filter(id=id,capacitacion_asistencia__asistio=True).values_list('capacitacion_asistencia__empleado__codigo','capacitacion_asistencia__empleado__nombre','capacitacion_asistencia__empleado__division__descripcion'))
        #insertando filas a la tabla
        contador_invitados=0
        for codigo,nombre,departamento in listado_convocados_capacitacion:
            contador_invitados+=1
            row_cells = convocados.add_row().cells
            row_cells[0].text = str(contador_invitados)
            row_cells[1].text = str(codigo)
            row_cells[2].text = str(nombre)
            row_cells[3].text = str(departamento)
        # parrafo.add_run().add_break()
        parrafo = archivo.add_paragraph("")
        parrafo = archivo.add_paragraph("Atentamente")
        parrafo = archivo.add_paragraph("Recursos Humanos")

        ######################################################################
        parrafo = archivo.add_paragraph("")
        parrafo.add_run().add_break(docx.text.run.WD_BREAK.PAGE)
        archivo.add_heading('(ENTREGAR SOLO ESTA HOJA AL OFICIAL DE CAPACITACION)',1)
        convocados_no_asistieron = archivo.add_table(rows=1, cols=5,style="Table Grid")
        #encabezados de tabla
        hdr_cells = convocados_no_asistieron.rows[0].cells
        hdr_cells[0].text = 'Nro.'
        hdr_cells[1].text = 'CODIGO DE EMPLEADO'
        hdr_cells[2].text = 'NOMBRE'
        hdr_cells[3].text = 'JUSTIFICAR'
        hdr_cells[4].text = 'FIRMA DEL JEFE INMEDIATO'
        
        listado_convocados_capacitacion_no_asistieron=list(capacitacion_evento_capacitacion.objects.filter(id=id,capacitacion_asistencia__asistio=False).values_list('capacitacion_asistencia__empleado__codigo','capacitacion_asistencia__empleado__nombre','capacitacion_asistencia__motivo_inasistencia__descripcion'))
        #insertando filas a la tabla
        contador_invitados=0
        for codigo,nombre,descripcion in listado_convocados_capacitacion_no_asistieron:
            contador_invitados+=1
            row_cells = convocados_no_asistieron.add_row().cells
            row_cells[0].text = str(contador_invitados)
            row_cells[1].text = str(codigo)
            row_cells[2].text = str(nombre)
            row_cells[3].text = str(descripcion)
        # parrafo.add_run().add_break()
        archivo.add_heading('EN CASO DE NO PODER ASISTIR, POR FAVOR JUSTIFIQUE LAS RAZONES DE SU INASISTENCIA: ',1)
        
        #######################################################################################################
        nombre_archivo_creado=str(settings.STATICFILES_DIRS[0]) + '/capacitacion_evaluacion/'

        nombre_archivo_creado=nombre_archivo_creado+' Convocatoria a Capacitacion de Personal '+campania+'.docx'
      
        
        archivo.save(nombre_archivo_creado)
        archivoword_codificado=''
        with open(nombre_archivo_creado, "rb") as archivoword:
            archivoword_codificado = base64.b64encode(archivoword.read()).decode('utf-8')
            
        
        return Response({"documento":str(archivoword_codificado)},status= status.HTTP_200_OK)



def achivo_capacitacion_F439(id,):
    archivo=Document()
    # archivo.add_paragraph('F-329 V3\tRight')
    section = archivo.sections[0]
    header = section.header
    unicode_character = "\u2611"
    paragraph = header.paragraphs[0]
    paragraph.text = "\t\tF-399v4"
    paragraph.style = archivo.styles["Header"]
    styles = archivo.styles
    styles['Heading 1'].font.color.rgb = docx.shared.RGBColor(0, 0, 0) #establece el color del titulo
    titulo=archivo.add_heading("CONVOCATORIA PARA CAPACITACIÓN",1)
    parrafo = archivo.add_paragraph(" ")
    titulo.alignment = 1
    hoy=datetime.now()
    hoy_formateado=  datetime.strftime(hoy, '%d-%m-%Y')
    solicitud_objeto = capacitacion_evento_capacitacion.objects.get(id=id) if capacitacion_evento_capacitacion.objects.filter(id=id) else None 
    if solicitud_objeto==None:
        return Response({"mensaje":"El evento no existe"},status=status.HTTP_404_NOT_FOUND)

    mes_evento= (capacitacion_evento_capacitacion.objects.filter(id=id).values('fecha_inicio__month'))[0]['fecha_inicio__month'] if capacitacion_evento_capacitacion.objects.filter(id=id).values('fecha_inicio__month')  else None
    mes_evento_digito= (capacitacion_evento_capacitacion.objects.filter(id=id).values('fecha_inicio__month'))[0]['fecha_inicio__month'] if capacitacion_evento_capacitacion.objects.filter(id=id).values('fecha_inicio__month')  else None
    anio_evento= (capacitacion_evento_capacitacion.objects.filter(id=id).values('fecha_inicio__year'))[0]['fecha_inicio__year'] if capacitacion_evento_capacitacion.objects.filter(id=id).values('fecha_inicio__year')  else None
    dia_evento= (capacitacion_evento_capacitacion.objects.filter(id=id).values('fecha_inicio__day'))[0]['fecha_inicio__day'] if capacitacion_evento_capacitacion.objects.filter(id=id).values('fecha_inicio__day')  else None
    meses= ['Enero','Febrero','Marzo','Abril','Mayo','Junio','Julio','Agosto','Septiembre','Octubre','Noviembre','Diciembre']
    campania= solicitud_objeto.campania.nombre_campania  if solicitud_objeto!=None else None
    impartido_por= solicitud_objeto.impartido_por if solicitud_objeto!=None else None
    if mes_evento!=None:
        indice=(mes_evento)-1
        mes_evento=meses[indice]

    archivo.add_heading(str(hoy_formateado),1)
    archivo.add_heading('Estimado/a Colaborador',1)
    parrafo = archivo.add_paragraph(" ")
    parrafo = archivo.add_paragraph("Nos complace informarle que han sido seleccionados para participar en la Capacitación "+str(campania)+' impartido por '+str(impartido_por) +', la cual se realizara: ' )
    archivo.add_heading('Bajo la Modalidad:',1)
    
    
    modalidades= capacitacion_modalidad.objects.filter().values_list('descripcion',flat=True)
    modalidad_evento=solicitud_objeto.modalidad.descripcion
    
    for modalidad in modalidades:
        # print('modalidad',modalidad)
        parrafo=archivo.add_paragraph(str(modalidad) +'      '+ str(unicode_character if modalidad==modalidad_evento else "☐"))
    parrafo = archivo.add_paragraph(" ")

    parrafo = archivo.add_paragraph('Día/Mes/Año: '+str(dia_evento)+'/'+str(mes_evento_digito)+'/'+str(anio_evento) )

    
    parrafo = archivo.add_paragraph("Listado del o los convocados")
    ##indicadores
    # archivo.add_heading("Indicadores:",1)
    convocados = archivo.add_table(rows=1, cols=4,style="Table Grid")
    #encabezados de tabla
    hdr_cells = convocados.rows[0].cells
    hdr_cells[0].text = 'Nro.'
    hdr_cells[1].text = 'CODIGO'
    hdr_cells[2].text = 'NOMBRE'
    hdr_cells[3].text = 'DEPARTAMENTO'
    
    listado_convocados_capacitacion=list(capacitacion_evento_capacitacion.objects.filter(id=id,capacitacion_asistencia__asistio=True).values_list('capacitacion_asistencia__empleado__codigo','capacitacion_asistencia__empleado__nombre','capacitacion_asistencia__empleado__division__descripcion'))
    #insertando filas a la tabla
    contador_invitados=0
    for codigo,nombre,departamento in listado_convocados_capacitacion:
        contador_invitados+=1
        row_cells = convocados.add_row().cells
        row_cells[0].text = str(contador_invitados)
        row_cells[1].text = str(codigo)
        row_cells[2].text = str(nombre)
        row_cells[3].text = str(departamento)
    # parrafo.add_run().add_break()
    parrafo = archivo.add_paragraph("")
    parrafo = archivo.add_paragraph("Atentamente")
    parrafo = archivo.add_paragraph("Recursos Humanos")

    ######################################################################
    parrafo = archivo.add_paragraph("")
    parrafo.add_run().add_break(docx.text.run.WD_BREAK.PAGE)
    archivo.add_heading('(ENTREGAR SOLO ESTA HOJA AL OFICIAL DE CAPACITACION)',1)
    convocados_no_asistieron = archivo.add_table(rows=1, cols=5,style="Table Grid")
    #encabezados de tabla
    hdr_cells = convocados_no_asistieron.rows[0].cells
    hdr_cells[0].text = 'Nro.'
    hdr_cells[1].text = 'CODIGO DE EMPLEADO'
    hdr_cells[2].text = 'NOMBRE'
    hdr_cells[3].text = 'JUSTIFICAR'
    hdr_cells[4].text = 'FIRMA DEL JEFE INMEDIATO'
    
    listado_convocados_capacitacion_no_asistieron=list(capacitacion_evento_capacitacion.objects.filter(id=id,capacitacion_asistencia__asistio=False).values_list('capacitacion_asistencia__empleado__codigo','capacitacion_asistencia__empleado__nombre','capacitacion_asistencia__motivo_inasistencia__descripcion'))
    #insertando filas a la tabla
    contador_invitados=0
    for codigo,nombre,descripcion in listado_convocados_capacitacion_no_asistieron:
        contador_invitados+=1
        row_cells = convocados_no_asistieron.add_row().cells
        row_cells[0].text = str(contador_invitados)
        row_cells[1].text = str(codigo)
        row_cells[2].text = str(nombre)
        row_cells[3].text = str(descripcion)
    # parrafo.add_run().add_break()
    archivo.add_heading('EN CASO DE NO PODER ASISTIR, POR FAVOR JUSTIFIQUE LAS RAZONES DE SU INASISTENCIA: ',1)
    
    #######################################################################################################
    nombre_archivo_creado=str(settings.STATICFILES_DIRS[0]) + '/capacitacion_evaluacion/'

    nombre_archivo_creado=nombre_archivo_creado+' Convocatoria a Capacitacion de Personal '+campania+'.docx'
    
    
    archivo.save(nombre_archivo_creado)
    # archivoword_codificado=''
    # with open(nombre_archivo_creado, "rb") as archivoword:
    #     archivoword_codificado = base64.b64encode(archivoword.read()).decode('utf-8')
        
    
    return nombre_archivo_creado
